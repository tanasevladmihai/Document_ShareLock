import hashlib
import io
import math

import numpy as np
import requests
import streamlit as st
from docx import Document
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer


OLLAMA_API_URL = "http://llama-server:8080/completion"
MODEL_URL = "http://llama-server:8080/completion"
#MAX_OUTPUT_TOKENS_INSIGHTFUL = 1536
#MAX_OUTPUT_TOKENS_PRECISE = 2048

MAX_OUTPUT_TOKENS = 2048
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_TOKEN_SIZE = 220
CHUNK_TOKEN_OVERLAP = 40
ANSWER_MODES = {
    "Precise": {
        "temperature": 0.5,
        "top_p": 0.9,
        "seed_chunks": 6,
        "neighbor_window": 0,
        "max_chunks": 6,
    },
    "Insightful": {
        "temperature": 0.8,
        "top_p": 0.95,
        "seed_chunks": 3,
        "neighbor_window": 1,
        "max_chunks": 6,
    },
}
SUMMARY_BATCH_CHUNKS = 6
SUMMARY_BATCH_TOKENS = 256
SUMMARY_FINAL_TOKENS = 512
MAX_SUMMARY_CHARS = 4500

st.set_page_config(page_title="LLM Chat Interface", page_icon=":speech_balloon:")


DOCUMENT_STATE_KEYS = [
    "document_fingerprint",
    "document_name",
    "document_chunks",
    "document_embeddings",
    "document_summary",
    "document_summary_error",
]


@st.cache_resource(show_spinner="Loading local embedding model...")
def load_embedding_model():
    return SentenceTransformer(EMBEDDING_MODEL_NAME)


def initialize_document_state():
    for key in DOCUMENT_STATE_KEYS:
        if key not in st.session_state:
            st.session_state[key] = None


def clear_document_state():
    for key in DOCUMENT_STATE_KEYS:
        st.session_state[key] = None


def read_txt_file(file_bytes):
    return file_bytes.decode("utf-8", errors="ignore")


def read_pdf_file(file_bytes):
    text = ""
    pdf_reader = PdfReader(io.BytesIO(file_bytes))

    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    return text


def read_docx_file(file_bytes):
    text = ""
    document = Document(io.BytesIO(file_bytes))

    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text += row_text + "\n"

    return text


def extract_file_text(uploaded_file):
    if uploaded_file is None:
        return ""

    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if file_name.endswith(".txt"):
        return read_txt_file(file_bytes)

    if file_name.endswith(".pdf"):
        return read_pdf_file(file_bytes)

    if file_name.endswith(".docx"):
        return read_docx_file(file_bytes)

    raise ValueError("Unsupported file type.")


def clean_extracted_text(text):
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line).strip()


def get_file_fingerprint(uploaded_file):
    digest = hashlib.sha256(uploaded_file.getvalue()).hexdigest()
    return f"{uploaded_file.name}:{digest}"


def chunk_text_by_tokens(text, tokenizer):
    token_ids = tokenizer.encode(text, add_special_tokens=False)
    if not token_ids:
        return []

    chunks = []
    step_size = max(1, CHUNK_TOKEN_SIZE - CHUNK_TOKEN_OVERLAP)
    start = 0

    while start < len(token_ids):
        end = min(start + CHUNK_TOKEN_SIZE, len(token_ids))
        chunk_token_ids = token_ids[start:end]
        chunk_text = tokenizer.decode(
            chunk_token_ids,
            skip_special_tokens=True,
            clean_up_tokenization_spaces=True,
        ).strip()

        if chunk_text:
            chunks.append(
                {
                    "index": len(chunks) + 1,
                    "text": chunk_text,
                    "token_count": len(chunk_token_ids),
                }
            )

        if end == len(token_ids):
            break

        start += step_size

    total_chunks = len(chunks)
    for chunk in chunks:
        chunk["total_chunks"] = total_chunks

    return chunks


def embed_texts(model, texts):
    if not texts:
        return np.empty((0, 0), dtype=np.float32)

    embeddings = model.encode(
        texts,
        normalize_embeddings=True,
        convert_to_numpy=True,
        show_progress_bar=False,
    )

    if embeddings.ndim == 1:
        embeddings = embeddings.reshape(1, -1)

    return embeddings.astype(np.float32)


def parse_model_response(response_json):
    if "content" in response_json:
        return response_json["content"].strip()

    choices_data = response_json.get("choices", [{}])
    if choices_data:
        return choices_data[0].get("text", "").strip()

    return ""


def wrap_model_prompt(user_content):
    return (
        f"<start_of_turn>user\n{user_content.strip()}<end_of_turn>\n"
        "<start_of_turn>model\n"
    )


def call_model_completion(prompt, n_predict, temperature=0.2, top_p=0.9):
    response = requests.post(
        MODEL_URL,
        json={
            "prompt": prompt,
            "stop": ["<end_of_turn>", "<start_of_turn>"],
            "n_predict": n_predict,
            "temperature": temperature,
            "top_p": top_p,
            "cache_prompt": True,
        },
        timeout=300,
    )
    response.raise_for_status()
    return parse_model_response(response.json())


def format_chunks_for_prompt(chunks, include_scores=False):
    formatted_chunks = []

    for chunk in chunks:
        score_text = ""
        if include_scores and "score" in chunk:
            score_text = f" similarity={chunk['score']:.3f}"

        formatted_chunks.append(
            f"[chunk {chunk['index']}/{chunk['total_chunks']}{score_text}]\n{chunk['text']}"
        )

    return "\n\n".join(formatted_chunks)


def iter_chunk_batches(chunks, batch_size):
    for batch_start in range(0, len(chunks), batch_size):
        yield chunks[batch_start:batch_start + batch_size]


def truncate_summary(summary):
    if len(summary) <= MAX_SUMMARY_CHARS:
        return summary

    return summary[:MAX_SUMMARY_CHARS].rsplit(" ", 1)[0].strip() + "..."


def summarize_chunk_batch(batch):
    batch_context = format_chunks_for_prompt(batch)
    summary_request = f"""
You are summarizing part of a document for later question answering.
Write a concise factual summary of the important entities, claims, numbers, dates, requirements, and decisions.
Do not add facts that are not present.

<document_chunks>
{batch_context}
</document_chunks>
"""
    return call_model_completion(
        wrap_model_prompt(summary_request),
        n_predict=SUMMARY_BATCH_TOKENS,
        temperature=0.2,
    )


def summarize_document(chunks):
    if not chunks:
        return ""

    batch_summaries = []
    total_batches = math.ceil(len(chunks) / SUMMARY_BATCH_CHUNKS)

    for batch_number, batch in enumerate(iter_chunk_batches(chunks, SUMMARY_BATCH_CHUNKS), start=1):
        summary = summarize_chunk_batch(batch)
        if summary:
            batch_summaries.append(f"Section {batch_number}/{total_batches}: {summary}")

    if not batch_summaries:
        return ""

    if len(batch_summaries) == 1:
        return truncate_summary(batch_summaries[0])

    combined_summaries = "\n\n".join(batch_summaries)
    final_summary_request = f"""
Create one whole-document summary from the section summaries below.
Keep it concise but preserve key facts, entities, numbers, dates, requirements, and decisions.
Do not add facts that are not present.

<section_summaries>
{combined_summaries}
</section_summaries>
"""
    final_summary = call_model_completion(
        wrap_model_prompt(final_summary_request),
        n_predict=SUMMARY_FINAL_TOKENS,
        temperature=0.2,
    )

    return truncate_summary(final_summary or combined_summaries)


def index_uploaded_document(uploaded_file):
    embedding_model = load_embedding_model()
    extracted_text = clean_extracted_text(extract_file_text(uploaded_file))

    if not extracted_text:
        raise ValueError("The uploaded document did not contain extractable text.")

    chunks = chunk_text_by_tokens(extracted_text, embedding_model.tokenizer)

    if not chunks:
        raise ValueError("The uploaded document was too short or could not be chunked.")

    embeddings = embed_texts(embedding_model, [chunk["text"] for chunk in chunks])
    summary = ""
    summary_error = None

    try:
        summary = summarize_document(chunks)
    except Exception as error:
        summary_error = f"Document summary could not be generated: {error}"

    st.session_state.document_fingerprint = get_file_fingerprint(uploaded_file)
    st.session_state.document_name = uploaded_file.name
    st.session_state.document_chunks = chunks
    st.session_state.document_embeddings = embeddings
    st.session_state.document_summary = summary
    st.session_state.document_summary_error = summary_error


def retrieve_relevant_chunks(query, mode_config):
    chunks = st.session_state.document_chunks or []
    embeddings = st.session_state.document_embeddings

    if not chunks or embeddings is None or embeddings.size == 0:
        return []

    embedding_model = load_embedding_model()
    query_embedding = embed_texts(embedding_model, [query])[0]
    scores = embeddings @ query_embedding
    seed_count = min(mode_config["seed_chunks"], len(chunks))
    ranked_seed_indices = np.argsort(scores)[::-1][:seed_count]
    neighbor_window = mode_config["neighbor_window"]

    if neighbor_window == 0:
        selected_indices = list(ranked_seed_indices)
    else:
        selected_indices = []
        seen_indices = set()

        for seed_index in ranked_seed_indices:
            seed_index = int(seed_index)
            expanded_start = max(0, seed_index - neighbor_window)
            expanded_end = min(len(chunks), seed_index + neighbor_window + 1)

            for chunk_index in range(expanded_start, expanded_end):
                if chunk_index not in seen_indices:
                    selected_indices.append(chunk_index)
                    seen_indices.add(chunk_index)

        selected_indices = selected_indices[:mode_config["max_chunks"]]
        selected_indices = sorted(selected_indices)

    retrieved_chunks = []
    for chunk_index in selected_indices:
        chunk = dict(chunks[int(chunk_index)])
        chunk["score"] = float(scores[int(chunk_index)])
        retrieved_chunks.append(chunk)

    return retrieved_chunks


def get_mode_instructions(answer_mode):
    if answer_mode == "Insightful":
        return """
You are an insightful document assistant. You may look for subtle patterns, risks, omissions, implications, and details someone might miss.
Stay grounded: separate document-supported observations from cautious interpretations.
If an idea is not directly established by the provided context, label it as an inference or a possibility.
Do not invent facts, names, dates, numbers, or requirements.
"""

    return """
You are a precise AI assistant. Answer the user's request using only the provided context.
If the answer is not supported by the provided context, say that the document does not provide enough information.
Do not invent facts or instructions.
"""


def build_direct_prompt(user_request, answer_mode):
    mode_instructions = get_mode_instructions(answer_mode)
    user_prompt_content = f"""
<instruction>
{mode_instructions}
There is no uploaded document context for this request, so answer directly and state uncertainty where appropriate.
</instruction>

<user_request>
{user_request}
</user_request>
"""
    return wrap_model_prompt(user_prompt_content)


def build_document_prompt(user_request, document_summary, retrieved_chunks, answer_mode):
    mode_instructions = get_mode_instructions(answer_mode)
    summary_text = document_summary or "No whole-document summary is available."
    chunk_context = format_chunks_for_prompt(retrieved_chunks, include_scores=True)

    user_prompt_content = f"""
<instruction>
{mode_instructions}
Answer the user's request using only the user request, the whole-document summary, and the retrieved document chunks below.
The summary gives broad document context. The chunks are the most relevant evidence.
</instruction>

<user_request>
{user_request}
</user_request>

<whole_document_summary>
{summary_text}
</whole_document_summary>

<retrieved_document_chunks>
{chunk_context}
</retrieved_document_chunks>
"""
    return wrap_model_prompt(user_prompt_content)

initialize_document_state()

st.title("Document ShareLock")
st.subheader("Chat Interface for Uploaded Documents")

st.write("Type a message, upload a document, or do both.")

user_message = st.text_area("Your message")

answer_mode = st.radio(
    "Answer mode",
    list(ANSWER_MODES.keys()),
    index=0,
    horizontal=True,
)

uploaded_file = st.file_uploader(
    "Upload a document",
    type=["txt", "pdf", "docx"]
)

if uploaded_file is not None:
    st.success(f"Uploaded file: {uploaded_file.name}")

    current_fingerprint = get_file_fingerprint(uploaded_file)
    if st.session_state.document_fingerprint != current_fingerprint:
        with st.spinner("Indexing document and preparing summary..."):
            try:
                index_uploaded_document(uploaded_file)
            except Exception as error:
                clear_document_state()
                st.session_state.document_summary_error = str(error)

    chunk_count = len(st.session_state.document_chunks or [])
    if chunk_count:
        summary_status = "summary ready" if st.session_state.document_summary else "summary unavailable"
        st.caption(f"Indexed {chunk_count} chunks; {summary_status}.")
        if st.session_state.document_summary_error:
            st.warning(st.session_state.document_summary_error)
    elif st.session_state.document_summary_error:
        st.warning(st.session_state.document_summary_error)
else:
    if st.session_state.document_fingerprint is not None:
        clear_document_state()

if st.button("Send"):
    user_request = user_message.strip()
    has_document = bool(st.session_state.document_chunks)
    mode_config = ANSWER_MODES[answer_mode]

    if user_request == "" and not has_document:
        st.warning("Please type a message or upload a document.")
    elif uploaded_file is not None and not has_document:
        st.warning("The uploaded document did not produce a searchable index.")
    else:
        if not user_request:
            user_request = "Summarize the uploaded document."

        if has_document:
            retrieved_chunks = retrieve_relevant_chunks(user_request, mode_config)
            final_prompt = build_document_prompt(
                user_request,
                st.session_state.document_summary,
                retrieved_chunks,
                answer_mode,
            )
            st.caption(f"{answer_mode} mode is using {len(retrieved_chunks)} retrieved chunks for this answer.")
        else:
            final_prompt = build_direct_prompt(user_request, answer_mode)

        with st.spinner("Sending to the model..."):
            try:
                generated_output = call_model_completion(
                    final_prompt,
                    n_predict=MAX_OUTPUT_TOKENS,
                    temperature=mode_config["temperature"],
                    top_p=mode_config["top_p"],
                )
                st.subheader("Response")

                if generated_output:
                    st.write(generated_output)
                else:
                    st.write("No message content found.")

            except Exception as error:
                st.error("Could not connect to the model server.")
                st.write(error)
